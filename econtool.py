# lost_earnings_app.py  –  Guided UI
# Year-by-Year Lost-Earnings (Past vs Future + Offsets)
# © 2025  Christopher Skerritt  – MIT License
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from datetime import date, timedelta
from dataclasses import dataclass
from typing import List, Sequence, Optional
import io, tempfile
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import RGBColor

# Import authentication module
from auth import check_authentication, get_current_user

# ════════════════════════════════════════════════════════════════
# 0.  Page configuration & authentication
# ════════════════════════════════════════════════════════════════
st.set_page_config(page_title="Lost-Earnings Schedule",
                   layout="wide", initial_sidebar_state="expanded")

# Check authentication first - this will show login form if not authenticated
if not check_authentication():
    st.stop()

# Get current user info
current_user = get_current_user()

st.title("Lost-Earnings Schedule Generator")
st.caption(
    "👉 **Purpose:** Quickly create *Past* and *Future* lost-earnings tables "
    "with optional mitigation offsets, Tinari adjustments, and present-value "
    "discounting.  Each section below contains helper text to guide you."
)

# Show welcome message with user name
st.success(f"Welcome back, {current_user['name']}! 👋")

# ════════════════════════════════════════════════════════════════
# 1.  Helper objects & core math                                  
# ════════════════════════════════════════════════════════════════
@dataclass
class Factor:
    label: str
    value: float        # decimal (e.g., 0.035 = 3.5 %)

def build_aif(factors: Sequence[Factor]) -> float:
    a = 1.0
    for f in factors:
        if f.label == "Personal Consumption":
            # Personal consumption uses (a - n) formula
            a *= (a - f.value)
        elif f.label == "Fringe Benefits":
            # Fringe benefits add to earnings: (1 + n) formula
            a *= (1 + f.value)
        else:
            # Other factors reduce earnings: (1 - n) formula
            a *= (1 - f.value)
    return round(a, 6)

def days_in_year(y): return 366 if (y % 4 == 0 and (y % 100 != 0 or y % 400 == 0)) else 365

def calculate_date_from_years(birth_date: date, years: float) -> date:
    """Calculate a date by adding years to a birth date"""
    try:
        # Add the whole years
        whole_years = int(years)
        fractional_year = years - whole_years

        # Calculate the target year
        target_year = birth_date.year + whole_years

        # Try to create the date with the same month and day
        try:
            base_date = date(target_year, birth_date.month, birth_date.day)
        except ValueError:
            # Handle leap year edge case (Feb 29)
            base_date = date(target_year, birth_date.month, 28)

        # Add the fractional year as days
        days_to_add = int(fractional_year * 365.25)
        result_date = base_date + timedelta(days=days_to_add)

        return result_date
    except:
        # Fallback to birth date if calculation fails
        return birth_date

# ════════════════════════════════════════════════════════════════
# 2.  Single-period schedule builder                              
# ════════════════════════════════════════════════════════════════
def schedule_block(
    dob: date, start: date, end_year: int,
    pre_base: float, pre_g: float,
    off_base: float, off_g: float,
    factors: Sequence[Factor],
    disc: float, pv_on: bool,
) -> pd.DataFrame:

    yrs = np.arange(start.year, end_year + 1)
    t   = np.arange(len(yrs))

    first_frac = (days_in_year(start.year) -
                 (start.timetuple().tm_yday - 1)) / days_in_year(start.year)
    portion = np.ones_like(t, float); portion[0] = first_frac

    ages = [round((date(y,7,1) - dob).days/365.25, 2) for y in yrs]
    pre  = pre_base  * (1+pre_g)**t
    off  = off_base  * (1+off_g)**t
    nom  = (pre - off)*portion
    aif  = build_aif(factors) if factors else 1.0
    adj  = nom*aif
    pv   = adj/((1+disc)**t) if pv_on else np.nan

    df = pd.DataFrame({
        "Calendar Year": yrs,
        "Portion of Year (%)": (portion*100).round(2),
        "Age (yrs)": ages,
        "Pre-Injury Earnings ($)": pre.round(2),
        "Mitigating/Offset Earnings ($)": (off*portion).round(2),
        "Nominal Loss ($)": nom.round(2),
        "AIF (%)": round(aif*100,2),
        "AIF-Adjusted Loss ($)": adj.round(2),
    })
    if pv_on: df["PV Loss ($)"] = pv.round(2)
    return df

# ════════════════════════════════════════════════════════════════
# 3.  SIDEBAR – Global parameters                                 
# ════════════════════════════════════════════════════════════════
with st.sidebar:
    st.header("🗓 Key Dates")
    dob = st.date_input("Date of birth", date(1980,1,1),
                        help="Used only for the Age column.",
                        format="MM/DD/YYYY")
    doi = st.date_input("Date of Injury (DOI)", date(2023,1,1),
                        help="Loss period starts the day after injury.",
                        format="MM/DD/YYYY")
    dor = st.date_input("Date of Report (DOR)", date(2025,6,1),
                        help="Splits Past vs Future; losses up to DOR are 'Past'.",
                        format="MM/DD/YYYY")

    st.markdown("**📊 Life Expectancy Calculations**")
    le = st.number_input("Life Expectancy (LE) in years",
                         0.0, 120.0, 78.5, step=0.1,
                         help="Total life expectancy from date of injury")
    wle = st.number_input("Worklife Expectancy (WLE) in years",
                          0.0, 80.0, 45.0, step=0.1,
                          help="Working life expectancy from date of injury")
    yfs = st.number_input("Years to Final Separation (YFS)",
                          0.0, 50.0, 0.0, step=0.1,
                          help="Additional years beyond retirement")

    # Calculate statistical dates
    statistical_death_date = calculate_date_from_years(doi, le)  # Changed from dob to doi
    statistical_retirement_date = calculate_date_from_years(doi, wle)
    ret_year = statistical_retirement_date.year

    # Display calculated dates
    st.markdown("**📅 Calculated Statistical Dates**")
    st.info(f"**Statistical Date of Death:** {statistical_death_date.strftime('%m/%d/%Y')}")
    st.info(f"**Statistical Retirement Date:** {statistical_retirement_date.strftime('%m/%d/%Y')} (Year: {ret_year})")

    # Work Life Ratio Calculator
    st.markdown("**⚖️ Work Life Ratio Calculator**")

    # Initialize work life factor
    if 'work_life_factor' not in st.session_state:
        st.session_state.work_life_factor = 0.91  # Default value

    if st.button("Calculate Work Life Ratio"):
        if yfs > 0:
            work_life_ratio = wle / yfs
            # Update the work life factor in session state
            st.session_state.work_life_factor = work_life_ratio

            if work_life_ratio >= 1:
                st.error(f"⚠️ Work Life Ratio = {work_life_ratio:.4f} (≥ 1.0)")
                st.warning("Please verify your WLE and YFS values. The ratio should be less than 1.0")
            else:
                st.success(f"✅ Work Life Ratio = {work_life_ratio:.4f}")
        else:
            st.warning("Please enter a value greater than 0 for Years to Final Separation (YFS)")

    # Display current work life factor
    st.info(f"Current WorkLife Factor: {st.session_state.work_life_factor:.4f}")

    st.markdown("---")
    st.header("💵 Wage Paths")
    pre_base  = st.number_input("Base *Pre-injury* earnings ($/yr)",
                                0.0, 1e7, 70_000.0, step=1_000.0,
                                format="%.2f",
                                help="Annualized but-for wages in DOI year.")
    pre_g     = st.number_input("Pre-injury growth rate (%/yr)",
                                0.0, 20.0, 3.0)/100

    st.caption("**Mitigation / Offset Earnings**")
    off_past_base = st.number_input("Offset base, Past period ($/yr)",
                                    0.0, 1e7, 0.0, step=1_000.0,
                                    format="%.2f",
                                    help="Set to 0 for straight loss in Past period.")
    off_past_g    = st.number_input("Offset growth, Past (%/yr)",
                                    0.0, 20.0, 0.0)/100
    off_fut_base  = st.number_input("Offset base, Future period ($/yr)",
                                    0.0, 1e7, 30_000.0, step=1_000.0,
                                    format="%.2f",
                                    help="Projected post-injury wages after report date.")
    off_fut_g     = st.number_input("Offset growth, Future (%/yr)",
                                    0.0, 20.0, 1.5)/100

    st.markdown("---")
    st.header("📉 Present-Value")
    pv_on = st.checkbox("Apply PV discount", True)
    disc  = st.number_input("Nominal discount rate (%/yr)",
                            0.0, 20.0, 4.0)/100 if pv_on else 0.0

    st.markdown("---")
    st.header("🪜 Tinari Adjustment Ladder")
    unemp  = st.number_input("Unemployment (%)", 0.0, 30.0, 3.5)/100
    tax    = st.number_input("Tax / offsets (%)", 0.0, 50.0, 12.0)/100

    # Fringe benefit deduction with editable percentage
    fringe_enabled = st.checkbox("Add fringe-benefit deduction")
    fringe_pct = 0.0
    if fringe_enabled:
        fringe_pct = st.number_input("Fringe-benefit deduction (%)", 0.0, 50.0, 6.0)/100

    # Wrongful death with editable percentage
    wrongful_enabled = st.checkbox("Wrongful-death (subtract personal consumption)")
    wrongful_pct = 0.0
    if wrongful_enabled:
        wrongful_pct = st.number_input("Personal consumption (%)", 0.0, 50.0, 25.0)/100

# Build factor list with explanatory ordering
factors: List[Factor] = []
if fringe_enabled: factors.append(Factor("Fringe Benefits", fringe_pct))
factors.append(Factor("Unemployment", unemp))
factors.append(Factor("Tax / Offsets", tax))
if wrongful_enabled: factors.append(Factor("Personal Consumption", wrongful_pct))

# Get work life factor and create complete factors list for calculations
worklife_factor = st.session_state.get('work_life_factor', 0.91)
complete_factors = factors.copy()
# Add work life factor as a reduction factor (1 - worklife_factor)
complete_factors.append(Factor("WorkLife Adjustment", 1 - worklife_factor))

# Calculate AEF and create detailed breakdown
base_aif = build_aif(factors)
aef_value = base_aif * worklife_factor

# Create Adjusted Earnings Factor table
st.markdown("---")
st.subheader("📊 Adjusted Earnings Factor (AEF)")

# Create the AEF breakdown table with formulas
aef_data = []
current_factor = 1.0

# Gross Earnings Base
aef_data.append(["Gross Earnings Base", "100.00%", "1.00", "Base = 1.0"])

# WorkLife Factor - use the calculated value from session state
worklife_factor = st.session_state.get('work_life_factor', 0.91)
current_factor *= worklife_factor
aef_data.append(["WorkLife Factor", f"{worklife_factor*100:.2f}%", f"{worklife_factor:.2f}", f"WLE/YFS = {worklife_factor:.4f}"])

# Apply each factor with formulas
for factor in factors:
    if factor.label == "Unemployment":
        current_factor *= (1 - factor.value)
        aef_data.append([f"Unemployment Factor", f"{factor.value*100:.2f}%", f"{factor.value:.2f}", f"(1 - {factor.value:.3f}) = {1-factor.value:.3f}"])
    elif factor.label == "Tax / Offsets":
        current_factor *= (1 - factor.value)
        aef_data.append([f"Tax Liability", f"{factor.value*100:.2f}%", f"{factor.value:.2f}", f"(1 - {factor.value:.3f}) = {1-factor.value:.3f}"])
    elif factor.label == "Personal Consumption":
        if factor.value > 0:
            # Personal consumption reduces from current accumulated factor
            reduction = current_factor - factor.value
            aef_data.append([f"Personal Consumption", f"{factor.value*100:.2f}%", f"{factor.value:.2f}", f"({current_factor:.3f} - {factor.value:.3f}) = {reduction:.3f}"])
            current_factor = reduction
    elif factor.label == "Fringe Benefits":
        current_factor *= (1 + factor.value)
        aef_data.append([f"Fringe Benefits", f"{factor.value*100:.2f}%", f"{factor.value:.2f}", f"(1 + {factor.value:.3f}) = {1+factor.value:.3f}"])

# Final AEF
aef_data.append(["", "", "", ""])
aef_data.append([f"Final AEF", f"{aef_value*100:.2f}%", f"{aef_value:.2f}", f"Total Calculation Result"])

# Display as DataFrame with formula column
aef_df = pd.DataFrame(aef_data, columns=["Factor", "Percentage", "Decimal", "Formula"])
st.dataframe(aef_df, use_container_width=True, hide_index=True)

# Add explanation
st.markdown("""
**📋 AEF Formula Explanations:**
- **WorkLife Factor**: Direct ratio of WLE/YFS
- **Unemployment**: Reduces earnings by unemployment rate: (1 - rate)
- **Tax Liability**: Reduces earnings by tax rate: (1 - rate)
- **Personal Consumption**: Subtracts from accumulated factor: (current - rate)
- **Fringe Benefits**: Adds to earnings: (1 + rate)
- **Final AEF**: Product of all factor adjustments
""")

# ════════════════════════════════════════════════════════════════
# 4.  Compute PAST and FUTURE tables                              
# ════════════════════════════════════════════════════════════════
df_past = schedule_block(
    dob, doi, dor.year,
    pre_base, pre_g,
    off_past_base, off_past_g,
    complete_factors, disc, pv_on,
)
# fraction past vs future in report year
past_days = (dor - date(dor.year,1,1)).days+1
fut_days  = days_in_year(dor.year) - past_days
past_frac, fut_frac = past_days/days_in_year(dor.year), fut_days/days_in_year(dor.year)
# Scale last row of past table
df_past.loc[df_past.index[-1], "Portion of Year (%)"] *= past_frac
for c in df_past.columns[3:]:   # money cols
    df_past.loc[df_past.index[-1], c] *= past_frac

df_future = schedule_block(
    dob, date(dor.year,1,1), ret_year,
    pre_base*(1+pre_g)**(dor.year-doi.year), pre_g,
    off_fut_base, off_fut_g,
    complete_factors, disc, pv_on,
)
df_future.loc[df_future.index[0], "Portion of Year (%)"] *= fut_frac
for c in df_future.columns[3:]:
    df_future.loc[df_future.index[0], c] *= fut_frac

# ════════════════════════════════════════════════════════════════
# 5.  Display results                                             
# ════════════════════════════════════════════════════════════════
st.markdown("---")
st.subheader("Past Losses  (DOI → DOR)")
st.caption("No mitigation shown if Past offset base is set to 0.")
st.dataframe(df_past, use_container_width=True)

st.subheader("Future Losses  (DOR → Retirement)")
st.caption("Offset column reflects claimant’s post-injury earnings path.")
st.dataframe(df_future, use_container_width=True)

# Totals summary
tot_cols = [c for c in df_past.columns if c.endswith("($)")]
summary_df = pd.DataFrame({
    "Past":   df_past[tot_cols].sum(),
    "Future": df_future[tot_cols].sum(),
}).T
st.subheader("Quick Totals")
st.dataframe(summary_df, use_container_width=True)

# ════════════════════════════════════════════════════════════════
# 6.  Downloads                                                   
# ════════════════════════════════════════════════════════════════
csv_past   = df_past.to_csv(index=False).encode()
csv_future = df_future.to_csv(index=False).encode()
excel_io   = io.BytesIO()
with pd.ExcelWriter(excel_io, engine="openpyxl") as writer:
    df_past.to_excel(writer, sheet_name="Past", index=False)
    df_future.to_excel(writer, sheet_name="Future", index=False)
excel_io.seek(0)

colc1, colc2, colc3 = st.columns(3)
colc1.download_button("⬇️ Past CSV", csv_past, "past_losses.csv")
colc2.download_button("⬇️ Future CSV", csv_future, "future_losses.csv")
colc3.download_button("⬇️ Excel (both)", excel_io,
                      "lost_earnings_split.xlsx",
                      "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ════════════════════════════════════════════════════════════════
# 7.  Charts                                                      
# ════════════════════════════════════════════════════════════════
st.markdown("---")
st.subheader("Charts")

# Earnings paths line chart
chart_df = df_future.set_index("Calendar Year")[[
    "Pre-Injury Earnings ($)",
    "Mitigating/Offset Earnings ($)",
    "Nominal Loss ($)",
]]
if pv_on:
    chart_df["PV Loss ($)"] = df_future["PV Loss ($)"]
fig1, ax1 = plt.subplots()
chart_df.plot(ax=ax1)
ax1.set_ylabel("Dollars")
ax1.set_title("Earnings Paths & Losses (Future period)")
fig1.tight_layout(); buf1 = io.BytesIO(); fig1.savefig(buf1, format="png", dpi=200)
st.pyplot(fig1)

# Past vs Future bar
fig2, ax2 = plt.subplots()
summary_df.plot(kind="barh", ax=ax2)
ax2.set_xlabel("Total Dollars")
ax2.set_title("Past vs Future – Cumulative Values")
fig2.tight_layout(); buf2 = io.BytesIO(); fig2.savefig(buf2, format="png", dpi=200)
st.pyplot(fig2)

# Tinari stacked bar (future)
buf3 = None
if "AIF-Adjusted Loss ($)" in df_future:
    bar_df = pd.DataFrame({
        "Calendar Year": df_future["Calendar Year"],
        "Nominal Loss": df_future["Nominal Loss ($)"],
        "AIF Deductions": df_future["Nominal Loss ($)"] - df_future["AIF-Adjusted Loss ($)"],
    }).set_index("Calendar Year")
    fig3, ax3 = plt.subplots()
    bar_df.plot(kind="bar", stacked=True, ax=ax3)
    ax3.set_ylabel("Dollars")
    ax3.set_title("Tinari Adjustments (Future)")
    fig3.tight_layout(); buf3 = io.BytesIO(); fig3.savefig(buf3, format="png", dpi=200)
    st.pyplot(fig3)

# ════════════════════════════════════════════════════════════════
# 8.  Word report                                                 
# ════════════════════════════════════════════════════════════════
def word_report(past_df, fut_df, charts, aef_df):
    doc = Document()

    # Set document to landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Swap page dimensions for landscape
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    # Add main heading and ensure black text
    main_heading = doc.add_heading("Lost-Earnings Analysis", 0)
    for run in main_heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Add key information section
    key_heading = doc.add_heading("Key Dates and Parameters", level=1)
    for run in key_heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = "Table Grid"  # Changed from "Light Shading Accent 1" to black and white

    # Add key dates and parameters
    info_data = [
        ["Date of Birth", dob.strftime('%m/%d/%Y')],
        ["Date of Injury", doi.strftime('%m/%d/%Y')],
        ["Date of Report", dor.strftime('%m/%d/%Y')],
        ["Life Expectancy (LE)", f"{le:.2f} years"],
        ["Worklife Expectancy (WLE)", f"{wle:.2f} years"],
        ["Statistical Death Date", statistical_death_date.strftime('%m/%d/%Y')],
        ["Statistical Retirement Date", statistical_retirement_date.strftime('%m/%d/%Y')],
        ["Final AEF", f"{aef_value:.2f}"]
    ]

    for i, (label, value) in enumerate(info_data):
        # Set label cell
        label_cell = info_table.cell(i, 0)
        label_cell.text = label
        # Ensure black text
        for paragraph in label_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Set value cell with proper formatting
        value_cell = info_table.cell(i, 1)
        if isinstance(value, (int, float)) and not pd.isna(value):
            # Check if this is the Final AEF value (should be formatted as percentage)
            if label == "Final AEF":
                value_cell.text = f"{value * 100:.2f}%"
            elif abs(value) >= 1000:
                value_cell.text = f"{value:,.2f}"
            else:
                value_cell.text = f"{value:.2f}"
        else:
            value_cell.text = str(value).replace("**", "")  # Remove any asterisks
        # Ensure black text
        for paragraph in value_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Add AEF table
    doc.add_page_break()
    aef_heading = doc.add_heading("Adjusted Earnings Factor (AEF)", level=1)
    for run in aef_heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    _tbl(doc, aef_df)

    # Add loss tables
    doc.add_page_break()
    past_heading = doc.add_heading("Past Losses (DOI → DOR)", level=1)
    for run in past_heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    _tbl(doc, past_df)

    doc.add_page_break()
    future_heading = doc.add_heading("Future Losses (DOR → Retirement)", level=1)
    for run in future_heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    _tbl(doc, fut_df)

    # Add charts
    titles = ["Earnings Paths (Future)", "Past vs Future Summary", "Tinari Adjustments"]
    for buf, title in zip(charts, titles):
        if buf:
            buf.seek(0)
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            tmp.write(buf.read())
            tmp.close()
            doc.add_page_break()
            chart_heading = doc.add_heading(title, level=2)
            for run in chart_heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            doc.add_picture(tmp.name, width=Inches(6))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()

def _tbl(doc, df):
    r, c = df.shape
    t = doc.add_table(r+1, c)
    t.style = "Table Grid"  # Changed from "Light Shading Accent 1" to black and white

    # Set column widths based on content type
    for i, col in enumerate(t.columns):
        for cell in col.cells:
            if i == 0:  # Factor name column
                cell.width = Inches(2.5)
            elif i == 3 and c > 3:  # Formula column if it exists
                cell.width = Inches(3.0)
            else:
                cell.width = Inches(1.2)

    # Add headers
    for j, col in enumerate(df.columns):
        cell = t.cell(0, j)
        cell.text = str(col)
        # Make header bold and black
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Ensure black text

    # Add data
    for i in range(r):
        for j in range(c):
            cell_value = df.iat[i, j]
            column_name = df.columns[j]

            # Format numeric values based on column type
            if isinstance(cell_value, (int, float)) and not pd.isna(cell_value):
                # Check if this is a percentage column
                if "%" in column_name:
                    # For percentage columns, format as XX.XX%
                    if "Portion of Year" in column_name or "AIF" in column_name:
                        # These are already multiplied by 100
                        cell_text = f"{cell_value:.2f}%"
                    else:
                        # These need to be multiplied by 100 (like decimal values in AEF table)
                        cell_text = f"{cell_value * 100:.2f}%" if cell_value <= 1 else f"{cell_value:.2f}%"
                elif abs(cell_value) >= 1000:
                    cell_text = f"{cell_value:,.2f}"
                else:
                    cell_text = f"{cell_value:.2f}"
            else:
                # Handle string values (like pre-formatted percentages in AEF table)
                cell_text = str(cell_value)
                # If it's in the Percentage column but already a string with %, ensure proper formatting
                if "Percentage" in column_name and "%" in cell_text:
                    # Extract the numeric part and reformat to ensure 2 decimal places
                    try:
                        numeric_part = float(cell_text.replace("%", ""))
                        cell_text = f"{numeric_part:.2f}%"
                    except ValueError:
                        # If conversion fails, keep original string
                        pass

            # Remove asterisks from text
            cell_text = cell_text.replace("**", "")

            cell = t.cell(i+1, j)
            cell.text = cell_text

            # Make final AEF row bold and black
            if "Final AEF" in cell_text:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Ensure black text
            else:
                # Ensure all other text is black
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)

doc_bytes = word_report(df_past, df_future, (buf1, buf2, buf3), aef_df)
st.download_button("⬇️ Word Report",
    doc_bytes, "lost_earnings_report.docx",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.caption("👆 Word report includes both tables and all charts.")
